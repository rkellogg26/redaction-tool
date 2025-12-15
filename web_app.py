#!/usr/bin/env python3
"""
Document Redaction Tool - Web Interface
"""

import os
import re
import uuid
import shutil
import tempfile
import zipfile
import threading
import time
import urllib.request
from pathlib import Path
from dataclasses import dataclass, field
from flask import Flask, render_template_string, request, send_file, jsonify, redirect, url_for
import fitz  # PyMuPDF
from docx import Document

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max for batch uploads

# Use temp directories for Render (ephemeral filesystem)
UPLOAD_FOLDER = Path(tempfile.gettempdir()) / 'redaction_uploads'
OUTPUT_FOLDER = Path(tempfile.gettempdir()) / 'redaction_outputs'
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)


@dataclass
class SensitiveCategory:
    name: str
    description: str
    patterns: list = field(default_factory=list)
    examples: list = field(default_factory=list)
    enabled: bool = True


PREDEFINED_CATEGORIES = {
    "ssn": SensitiveCategory(
        name="Social Security Numbers",
        description="XXX-XX-XXXX format",
        patterns=[r'\b\d{3}-\d{2}-\d{4}\b', r'\b\d{9}\b'],
        examples=["123-45-6789"]
    ),
    "email": SensitiveCategory(
        name="Email Addresses",
        description="user@domain.com",
        patterns=[r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'],
        examples=["example@email.com"]
    ),
    "phone": SensitiveCategory(
        name="Phone Numbers",
        description="Various US formats",
        patterns=[
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}\b'
        ],
        examples=["555-123-4567"]
    ),
    "creditcard": SensitiveCategory(
        name="Credit Card Numbers",
        description="13-19 digit card numbers",
        patterns=[r'\b(?:\d{4}[-\s]?){3,4}\d{1,4}\b'],
        examples=["4111-1111-1111-1111"]
    ),
    "date": SensitiveCategory(
        name="Dates",
        description="Various date formats",
        patterns=[
            r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',
            r'\b\d{1,2}-\d{1,2}-\d{2,4}\b',
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b'
        ],
        examples=["01/15/2024"]
    ),
    "corporate": SensitiveCategory(
        name="Corporate Buzzwords",
        description="Icky corporate jargon",
        patterns=[
            r'\bsynerg(?:y|ies|ize|izing|istic)\b',
            r'\bleverage[ds]?\b',
            r'\bparadigm(?:s)?\s*(?:shift)?\b',
            r'\bthought\s+leader(?:ship)?\b',
            r'\bbest[- ]in[- ]class\b',
            r'\bworld[- ]class\b',
            r'\bcut(?:ting)?[- ]edge\b',
            r'\bgame[- ]chang(?:er|ing)\b',
            r'\bmove(?:ing)?\s+the\s+needle\b',
            r'\blow[- ]hanging\s+fruit\b',
            r'\bcircle\s+back\b',
            r'\btouch\s+base\b',
            r'\bdeep\s+dive\b',
            r'\bdrill\s+down\b',
            r'\b(?:strategic|passionate|driven|dynamic|proactive|innovative|visionary)\b',
            r'\bstakeholder(?:s)?\b',
            r'\bvalue[- ]add(?:ed)?\b',
            r'\bwin[- ]win\b',
            r'\bscalable\b',
            r'\bagile\b',
            r'\bdisrupt(?:ive|ion|ing|or)?\b',
            r'\bimpact(?:ful)?\b',
            r'\brobust\b',
            r'\bholistic\b',
            r'\becosystem\b',
            r'\boptimize[ds]?\b',
            r'\bstreamline[ds]?\b',
            r'\bactionable(?:\s+insights?)?\b',
            r'\bbandwidth\b',
            r'\bcore\s+competenc(?:y|ies)\b',
            r'\bempower(?:ed|ing|ment)?\b',
            r'\bfacilitat(?:e[ds]?|ing|ion)\b',
            r'\bgranular(?:ity)?\b',
            r'\bincentivize[ds]?\b',
            r'\bintegrat(?:e[ds]?|ion)\b',
            r'\biterati(?:ve|on)\b',
            r'\bKPI(?:s)?\b',
            r'\bmetric(?:s)?\b',
            r'\bmilestone(?:s)?\b',
            r'\bmonetize[ds]?\b',
            r'\bonboard(?:ed|ing)?\b',
            r'\bpain\s+point(?:s)?\b',
            r'\bpipeline\b',
            r'\bpivot(?:ed|ing)?\b',
            r'\bproactive(?:ly)?\b',
            r'\broadmap\b',
            r'\bROI\b',
            r'\bscope\s+creep\b',
            r'\bsiloed?\b',
            r'\bsprint(?:s)?\b',
            r'\bsync(?:ed)?\s+up\b',
            r'\btake(?:away)?s?\b',
            r'\bteam\s+player\b',
            r'\bunpack(?:ing)?\b',
            r'\bvertical(?:s)?\b',
            r'\bwear(?:ing)?\s+many\s+hats\b',
            r'\bhit\s+the\s+ground\s+running\b',
            r'\bout\s+of\s+the\s+box\b',
            r'\bpush\s+the\s+envelope\b',
            r'\braise\s+the\s+bar\b',
            r'\bresults[- ](?:driven|oriented)\b',
            r'\bself[- ](?:starter|motivated)\b',
            r'\bskill\s*set\b',
            r'\bsolution(?:s)?[- ]oriented\b',
            r'\bthink\s+outside\s+the\s+box\b',
            r'\btrack\s+record\b',
        ],
        examples=["synergy"]
    ),
}


class SensitiveInfoDetector:
    def __init__(self):
        self.categories = {}
        self.custom_terms = []

    def add_category(self, key: str, category: SensitiveCategory):
        self.categories[key] = category

    def set_custom_terms(self, terms: list):
        self.custom_terms = [t.strip() for t in terms if t.strip()]

    def find_sensitive_text(self, text: str) -> list:
        matches = []

        for key, category in self.categories.items():
            if not category.enabled:
                continue
            for pattern in category.patterns:
                try:
                    for match in re.finditer(pattern, text, re.IGNORECASE):
                        matches.append((match.start(), match.end(), match.group(), category.name))
                except re.error:
                    continue

        for term in self.custom_terms:
            escaped = re.escape(term)
            try:
                for match in re.finditer(escaped, text, re.IGNORECASE):
                    matches.append((match.start(), match.end(), match.group(), "Custom Terms"))
            except re.error:
                continue

        matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))
        filtered = []
        last_end = -1
        for match in matches:
            if match[0] >= last_end:
                filtered.append(match)
                last_end = match[1]
        return filtered


def redact_pdf(input_path: str, output_path: str, detector: SensitiveInfoDetector) -> dict:
    doc = fitz.open(input_path)
    stats = {"pages": 0, "redactions": 0, "categories": {}}

    for page in doc:
        stats["pages"] += 1
        text_dict = page.get_text("dict")

        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text:
                        continue
                    matches = detector.find_sensitive_text(text)
                    for start, end, matched_text, category in matches:
                        instances = page.search_for(matched_text)
                        for inst in instances:
                            page.add_redact_annot(inst, fill=(0, 0, 0))
                            stats["redactions"] += 1
                            stats["categories"][category] = stats["categories"].get(category, 0) + 1
        page.apply_redactions()

    doc.save(output_path)
    doc.close()
    return stats


def redact_docx(input_path: str, output_path: str, detector: SensitiveInfoDetector) -> dict:
    doc = Document(input_path)
    stats = {"paragraphs": 0, "redactions": 0, "categories": {}}

    def process_paragraph(para):
        full_text = para.text
        if not full_text:
            return
        matches = detector.find_sensitive_text(full_text)
        if not matches:
            return

        new_text = list(full_text)
        for start, end, matched_text, category in matches:
            for i in range(start, end):
                new_text[i] = "█"
            stats["redactions"] += 1
            stats["categories"][category] = stats["categories"].get(category, 0) + 1

        new_text = "".join(new_text)
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new_text
            para.runs[0].font.name = font_name
            para.runs[0].font.size = font_size
            para.runs[0].font.bold = bold
            para.runs[0].font.italic = italic

    for para in doc.paragraphs:
        stats["paragraphs"] += 1
        process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    stats["paragraphs"] += 1
                    process_paragraph(para)

    doc.save(output_path)
    return stats


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text
    elif ext == ".docx":
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""


HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Redact | Document Privacy Tool</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700;900&display=swap" rel="stylesheet">
    <style>
        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }
        body {
            font-family: 'Roboto', -apple-system, BlinkMacSystemFont, sans-serif;
            background: #ffffff;
            min-height: 100vh;
            color: #fff;
            overflow-x: hidden;
        }
        body::before {
            content: '';
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background:
                radial-gradient(ellipse at 20% 20%, rgba(0, 48, 135, 0.08) 0%, transparent 50%),
                radial-gradient(ellipse at 80% 80%, rgba(0, 73, 171, 0.06) 0%, transparent 50%),
                radial-gradient(ellipse at 50% 50%, rgba(0, 48, 135, 0.04) 0%, transparent 70%);
            pointer-events: none;
            z-index: 0;
        }
        .container {
            max-width: 960px;
            margin: 0 auto;
            padding: 40px 24px;
            position: relative;
            z-index: 1;
        }
        .header {
            text-align: center;
            margin-bottom: 48px;
        }
        .logo {
            display: inline-flex;
            align-items: center;
            gap: 16px;
            margin-bottom: 16px;
        }
        .logo-text {
            font-family: 'Roboto', sans-serif;
            font-size: 2.2rem;
            font-weight: 800;
            letter-spacing: 0.15em;
            color: #000;
            display: flex;
            align-items: center;
        }
        .logo-text .redacted {
            display: inline-block;
            position: relative;
            width: 1.0em;
            height: 1.1em;
            margin: 0 2px;
            vertical-align: middle;
        }
        .logo-text .redacted .letter {
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            color: #000;
            z-index: 1;
        }
        .logo-text .redacted .bar {
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 0%;
            background: #000;
            border-radius: 3px;
            z-index: 2;
            animation: redact-drape 3s ease-out forwards;
            animation-delay: var(--delay, 0s);
        }
        .logo-text .redacted:nth-child(2) .bar {
            --delay: 1s;
        }
        .logo-text .redacted:nth-child(4) .bar {
            --delay: 2s;
        }
        @keyframes redact-drape {
            0% { height: 0%; }
            100% { height: 100%; }
        }
        h1 {
            font-size: 2.75rem;
            font-weight: 700;
            letter-spacing: -0.02em;
            background: linear-gradient(135deg, #ffffff 0%, #a8c5e8 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        .subtitle {
            color: #555;
            font-size: 1.1rem;
            font-weight: 400;
            max-width: 500px;
            margin: 0 auto;
            line-height: 1.6;
        }
        .card {
            background: #f8f9fa;
            border-radius: 20px;
            padding: 28px;
            margin-bottom: 24px;
            border: 1px solid #e0e0e0;
            transition: all 0.3s ease;
        }
        .card:hover {
            background: #f0f2f5;
            border-color: #ccc;
        }
        .card h2 {
            font-size: 0.85rem;
            font-weight: 600;
            margin-bottom: 20px;
            color: #333;
            text-transform: uppercase;
            letter-spacing: 0.1em;
        }
        .upload-zone {
            border: 2px dashed #ccc;
            border-radius: 16px;
            padding: 48px 32px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s ease;
            background: #fafafa;
        }
        .upload-zone:hover {
            border-color: #0073E6;
            background: rgba(0, 115, 230, 0.05);
            transform: translateY(-2px);
        }
        .upload-zone.dragover {
            border-color: #0073E6;
            background: rgba(0, 115, 230, 0.1);
            box-shadow: 0 0 40px rgba(0, 115, 230, 0.15);
        }
        .upload-icon {
            font-size: 56px;
            margin-bottom: 16px;
        }
        .upload-zone p {
            font-size: 1rem;
            color: #333;
        }
        .upload-zone .hint {
            color: #888;
            font-size: 0.9rem;
            margin-top: 8px;
        }
        .file-input {
            display: none;
        }
        .file-list {
            margin-top: 20px;
            display: none;
        }
        .file-list.visible {
            display: block;
        }
        .file-item {
            padding: 14px 18px;
            background: rgba(0, 115, 230, 0.08);
            border: 1px solid rgba(0, 115, 230, 0.2);
            border-radius: 12px;
            margin-bottom: 10px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            transition: all 0.2s ease;
        }
        .file-item:hover {
            background: rgba(0, 115, 230, 0.12);
        }
        .file-item span {
            font-size: 0.95rem;
            color: #333;
        }
        .file-item .remove-btn {
            background: rgba(220, 53, 69, 0.1);
            border: 1px solid rgba(220, 53, 69, 0.3);
            color: #dc3545;
            padding: 6px 14px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 0.85rem;
            font-weight: 500;
            transition: all 0.2s ease;
        }
        .file-item .remove-btn:hover {
            background: rgba(220, 53, 69, 0.2);
        }
        .categories {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(220px, 1fr));
            gap: 14px;
        }
        .category-item {
            display: flex;
            align-items: flex-start;
            padding: 16px;
            background: #fff;
            border: 1px solid #ddd;
            border-radius: 14px;
            cursor: pointer;
            transition: all 0.2s ease;
        }
        .category-item:hover {
            background: #f5f5f5;
            border-color: #bbb;
            transform: translateY(-1px);
        }
        .category-item.corporate-item {
            background: rgba(255, 152, 0, 0.08);
            border-color: rgba(255, 152, 0, 0.3);
        }
        .category-item.corporate-item:hover {
            background: rgba(255, 152, 0, 0.15);
            border-color: rgba(255, 152, 0, 0.5);
        }
        .category-item input[type="checkbox"] {
            width: 20px;
            height: 20px;
            margin-right: 14px;
            margin-top: 2px;
            accent-color: #0073E6;
            flex-shrink: 0;
        }
        .category-item label {
            cursor: pointer;
            flex: 1;
        }
        .category-name {
            font-weight: 600;
            font-size: 0.95rem;
            margin-bottom: 4px;
            color: #222;
        }
        .category-example {
            font-size: 0.8rem;
            color: #777;
        }
        .custom-terms textarea {
            width: 100%;
            height: 130px;
            background: #fff;
            border: 1px solid #ddd;
            border-radius: 14px;
            padding: 16px;
            color: #333;
            font-family: 'SF Mono', 'Fira Code', monospace;
            font-size: 14px;
            resize: vertical;
            transition: all 0.2s ease;
        }
        .custom-terms textarea:focus {
            outline: none;
            border-color: #0073E6;
            box-shadow: 0 0 0 3px rgba(0, 115, 230, 0.15);
        }
        .custom-terms textarea::placeholder {
            color: #999;
        }
        .custom-terms .hint {
            color: #666;
            font-size: 0.85rem;
            margin-bottom: 12px;
        }
        .buttons {
            display: flex;
            gap: 14px;
            flex-wrap: wrap;
        }
        .btn {
            padding: 16px 32px;
            border: none;
            border-radius: 12px;
            font-size: 1rem;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            display: inline-flex;
            align-items: center;
            gap: 10px;
            font-family: 'Roboto', sans-serif;
        }
        .btn-primary {
            background: linear-gradient(135deg, #0073E6 0%, #0049AB 100%);
            color: #fff;
            box-shadow: 0 4px 20px rgba(0, 115, 230, 0.4);
        }
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 30px rgba(0, 115, 230, 0.5);
        }
        .btn-secondary {
            background: rgba(255, 255, 255, 0.08);
            color: #fff;
            border: 1px solid rgba(255, 255, 255, 0.15);
        }
        .btn-secondary:hover {
            background: rgba(255, 255, 255, 0.12);
            border-color: rgba(255, 255, 255, 0.25);
        }
        .btn:disabled {
            opacity: 0.4;
            cursor: not-allowed;
            transform: none;
            box-shadow: none;
        }
        .results {
            display: none;
        }
        .results.visible {
            display: block;
            animation: fadeIn 0.4s ease;
        }
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }
        .results-header {
            display: flex;
            align-items: center;
            gap: 16px;
            margin-bottom: 24px;
        }
        .results-header .icon {
            width: 56px;
            height: 56px;
            background: linear-gradient(135deg, #00C853 0%, #00A844 100%);
            border-radius: 16px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 28px;
            box-shadow: 0 8px 24px rgba(0, 200, 83, 0.3);
        }
        .results-header h2 {
            font-size: 1.5rem;
            font-weight: 600;
            text-transform: none;
            letter-spacing: normal;
            color: #222;
        }
        .stats {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(140px, 1fr));
            gap: 14px;
            margin-bottom: 24px;
        }
        .stat-item {
            background: rgba(0, 115, 230, 0.08);
            border: 1px solid rgba(0, 115, 230, 0.2);
            padding: 20px;
            border-radius: 14px;
            text-align: center;
        }
        .stat-value {
            font-size: 2.25rem;
            font-weight: 700;
            color: #0073E6;
            line-height: 1;
            margin-bottom: 6px;
        }
        .stat-label {
            font-size: 0.8rem;
            color: #555;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }
        .preview-box {
            background: #fff;
            border: 1px solid #ddd;
            border-radius: 14px;
            padding: 20px;
            max-height: 320px;
            overflow-y: auto;
            font-family: 'SF Mono', 'Fira Code', monospace;
            font-size: 13px;
            line-height: 1.7;
            white-space: pre-wrap;
            word-break: break-word;
            color: #333;
        }
        .preview-box .match {
            background: linear-gradient(135deg, #FFD54F 0%, #FFB300 100%);
            color: #000;
            padding: 2px 6px;
            border-radius: 4px;
            font-weight: 600;
        }
        .loading {
            display: none;
            text-align: center;
            padding: 48px;
        }
        .loading.visible {
            display: block;
        }
        .spinner {
            width: 56px;
            height: 56px;
            border: 3px solid #e0e0e0;
            border-top-color: #0073E6;
            border-radius: 50%;
            animation: spin 0.8s linear infinite;
            margin: 0 auto 24px;
        }
        @keyframes spin {
            to { transform: rotate(360deg); }
        }
        .download-btn {
            display: none;
            margin-top: 24px;
        }
        .download-btn.visible {
            display: inline-flex;
        }
        .batch-results {
            margin-top: 20px;
        }
        .batch-item {
            background: #fff;
            border: 1px solid #ddd;
            padding: 16px 20px;
            border-radius: 12px;
            margin-bottom: 10px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .batch-item .file-info {
            flex: 1;
        }
        .batch-item .file-name {
            font-weight: 600;
            font-size: 0.95rem;
            color: #222;
        }
        .batch-item .file-stats {
            font-size: 0.85rem;
            color: #666;
            margin-top: 4px;
        }
        .progress-bar {
            width: 100%;
            height: 6px;
            background: #e0e0e0;
            border-radius: 3px;
            margin: 20px 0;
            overflow: hidden;
        }
        .progress-fill {
            height: 100%;
            background: linear-gradient(90deg, #0073E6 0%, #00A8E8 100%);
            border-radius: 3px;
            transition: width 0.3s ease;
            width: 0%;
        }
        .progress-text {
            text-align: center;
            margin-bottom: 12px;
            color: #555;
            font-size: 0.95rem;
        }
        .footer {
            text-align: center;
            margin-top: 48px;
            padding-top: 24px;
            border-top: 1px solid #e0e0e0;
            color: #999;
            font-size: 0.85rem;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="logo">
                <div class="logo-text">R<span class="redacted"><span class="letter">E</span><span class="bar"></span></span>D<span class="redacted"><span class="letter">A</span><span class="bar"></span></span>CT</div>
            </div>
            <p class="subtitle">Protect sensitive information in your documents with intelligent, automated redaction</p>
        </div>

        <form id="redactionForm" enctype="multipart/form-data">
            <div class="card">
                <h2>Upload Documents</h2>
                <div class="upload-zone" id="uploadZone">
                    <div class="upload-icon">📄</div>
                    <p>Drag & drop your files here</p>
                    <p class="hint">PDF and Word documents supported • Multiple files allowed</p>
                    <input type="file" name="files" id="fileInput" class="file-input" accept=".pdf,.docx" multiple>
                </div>
                <div class="file-list" id="fileList"></div>
            </div>

            <div class="card">
                <h2>What to Redact</h2>
                <div class="categories">
                    {% for key, cat in categories.items() %}
                    <div class="category-item{% if key == 'corporate' %} corporate-item{% endif %}">
                        <input type="checkbox" name="categories" value="{{ key }}" id="cat_{{ key }}" {% if key != 'corporate' %}checked{% endif %}>
                        <label for="cat_{{ key }}">
                            <div class="category-name">{{ cat.name }}</div>
                            <div class="category-example">{{ cat.description }}</div>
                        </label>
                    </div>
                    {% endfor %}
                </div>
            </div>

            <div class="card">
                <h2>Custom Terms</h2>
                <div class="custom-terms">
                    <p class="hint">Add specific names, phrases, or terms to redact (one per line)</p>
                    <textarea name="custom_terms" id="customTerms" placeholder="John Doe&#10;Project Phoenix&#10;Confidential"></textarea>
                </div>
            </div>

            <div class="card">
                <div class="buttons">
                    <button type="button" class="btn btn-secondary" id="previewBtn" disabled>
                        👁️ Preview
                    </button>
                    <button type="submit" class="btn btn-primary" id="redactBtn" disabled>
                        🔒 Redact Documents
                    </button>
                </div>
            </div>
        </form>

        <div class="card loading" id="loading">
            <div class="spinner"></div>
            <p class="progress-text" id="progressText">Processing your documents...</p>
            <div class="progress-bar">
                <div class="progress-fill" id="progressFill"></div>
            </div>
        </div>

        <div class="card results" id="results">
            <div class="results-header">
                <div class="icon">✓</div>
                <div>
                    <h2>Redaction Complete</h2>
                    <p style="color: rgba(255,255,255,0.5); margin-top: 4px; font-size: 0.95rem;" id="outputSummary"></p>
                </div>
            </div>
            <div class="stats" id="statsContainer"></div>
            <div class="batch-results" id="batchResults"></div>
            <a href="#" class="btn btn-primary download-btn" id="downloadBtn">
                📥 Download Redacted Documents
            </a>
        </div>

        <div class="card results" id="previewResults">
            <h2>Preview - Items to Redact</h2>
            <div class="stats" id="previewStats"></div>
            <h3 style="margin: 16px 0 8px; font-size: 1rem;">Document Preview:</h3>
            <div class="preview-box" id="previewContent"></div>
        </div>
    </div>

    <script>
        const uploadZone = document.getElementById('uploadZone');
        const fileInput = document.getElementById('fileInput');
        const fileList = document.getElementById('fileList');
        const previewBtn = document.getElementById('previewBtn');
        const redactBtn = document.getElementById('redactBtn');
        const form = document.getElementById('redactionForm');
        const loading = document.getElementById('loading');
        const results = document.getElementById('results');
        const previewResults = document.getElementById('previewResults');

        let selectedFiles = [];

        // File upload handling
        uploadZone.addEventListener('click', () => fileInput.click());

        uploadZone.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadZone.classList.add('dragover');
        });

        uploadZone.addEventListener('dragleave', () => {
            uploadZone.classList.remove('dragover');
        });

        uploadZone.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadZone.classList.remove('dragover');
            if (e.dataTransfer.files.length) {
                addFiles(e.dataTransfer.files);
            }
        });

        fileInput.addEventListener('change', () => {
            if (fileInput.files.length) {
                addFiles(fileInput.files);
            }
        });

        function addFiles(files) {
            for (const file of files) {
                const ext = file.name.split('.').pop().toLowerCase();
                if (['pdf', 'docx'].includes(ext)) {
                    // Check for duplicates
                    if (!selectedFiles.some(f => f.name === file.name && f.size === file.size)) {
                        selectedFiles.push(file);
                    }
                }
            }
            updateFileList();
        }

        function removeFile(index) {
            selectedFiles.splice(index, 1);
            updateFileList();
        }

        function updateFileList() {
            if (selectedFiles.length === 0) {
                fileList.classList.remove('visible');
                previewBtn.disabled = true;
                redactBtn.disabled = true;
                return;
            }

            fileList.classList.add('visible');
            previewBtn.disabled = false;
            redactBtn.disabled = false;

            fileList.innerHTML = selectedFiles.map((file, index) => `
                <div class="file-item">
                    <span>📎 ${file.name} (${(file.size / 1024).toFixed(1)} KB)</span>
                    <button type="button" class="remove-btn" onclick="removeFile(${index})">✕</button>
                </div>
            `).join('');
        }

        // Preview (first file only)
        previewBtn.addEventListener('click', async () => {
            if (selectedFiles.length === 0) return;

            const formData = new FormData();
            formData.append('file', selectedFiles[0]);

            document.querySelectorAll('input[name="categories"]:checked').forEach(cb => {
                formData.append('categories', cb.value);
            });
            formData.append('custom_terms', document.getElementById('customTerms').value);

            loading.classList.add('visible');
            document.getElementById('progressText').textContent = 'Analyzing document...';
            document.getElementById('progressFill').style.width = '50%';
            results.classList.remove('visible');
            previewResults.classList.remove('visible');

            try {
                const response = await fetch('/preview', {
                    method: 'POST',
                    body: formData
                });

                const data = await response.json();
                loading.classList.remove('visible');

                if (data.error) {
                    alert('Error: ' + data.error);
                    return;
                }

                let statsHtml = '';
                statsHtml += `<div class="stat-item"><div class="stat-value">${data.total}</div><div class="stat-label">Total Items</div></div>`;
                for (const [cat, count] of Object.entries(data.categories)) {
                    statsHtml += `<div class="stat-item"><div class="stat-value">${count}</div><div class="stat-label">${cat}</div></div>`;
                }
                document.getElementById('previewStats').innerHTML = statsHtml;
                document.getElementById('previewContent').innerHTML = data.preview_html;
                previewResults.classList.add('visible');

            } catch (err) {
                loading.classList.remove('visible');
                alert('Error: ' + err.message);
            }
        });

        // Batch Redact
        form.addEventListener('submit', async (e) => {
            e.preventDefault();

            if (selectedFiles.length === 0) return;

            loading.classList.add('visible');
            results.classList.remove('visible');
            previewResults.classList.remove('visible');

            const formData = new FormData();
            selectedFiles.forEach(file => {
                formData.append('files', file);
            });

            document.querySelectorAll('input[name="categories"]:checked').forEach(cb => {
                formData.append('categories', cb.value);
            });
            formData.append('custom_terms', document.getElementById('customTerms').value);

            document.getElementById('progressText').textContent = `Processing ${selectedFiles.length} document(s)...`;
            document.getElementById('progressFill').style.width = '30%';

            try {
                const response = await fetch('/redact-batch', {
                    method: 'POST',
                    body: formData
                });

                document.getElementById('progressFill').style.width = '90%';

                const data = await response.json();
                loading.classList.remove('visible');

                if (data.error) {
                    alert('Error: ' + data.error);
                    return;
                }

                // Show results
                document.getElementById('outputSummary').textContent =
                    `${data.file_count} document(s) processed, ${data.total_redactions} total redactions`;

                let statsHtml = '';
                statsHtml += `<div class="stat-item"><div class="stat-value">${data.file_count}</div><div class="stat-label">Documents</div></div>`;
                statsHtml += `<div class="stat-item"><div class="stat-value">${data.total_redactions}</div><div class="stat-label">Total Redactions</div></div>`;
                for (const [cat, count] of Object.entries(data.total_categories)) {
                    statsHtml += `<div class="stat-item"><div class="stat-value">${count}</div><div class="stat-label">${cat}</div></div>`;
                }
                document.getElementById('statsContainer').innerHTML = statsHtml;

                // Show individual file results
                let batchHtml = '';
                for (const file of data.files) {
                    batchHtml += `
                        <div class="batch-item">
                            <div class="file-info">
                                <div class="file-name">📄 ${file.original_name}</div>
                                <div class="file-stats">${file.redactions} redactions</div>
                            </div>
                        </div>
                    `;
                }
                document.getElementById('batchResults').innerHTML = batchHtml;

                const downloadBtn = document.getElementById('downloadBtn');
                downloadBtn.href = '/download/' + data.download_id;
                downloadBtn.textContent = data.file_count > 1 ? '📥 Download All (ZIP)' : '📥 Download Redacted Document';
                downloadBtn.classList.add('visible');

                results.classList.add('visible');

            } catch (err) {
                loading.classList.remove('visible');
                alert('Error: ' + err.message);
            }
        });
    </script>
</body>
</html>
'''


@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE, categories=PREDEFINED_CATEGORIES)


@app.route('/preview', methods=['POST'])
def preview():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'})

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'})

    ext = Path(file.filename).suffix.lower()
    if ext not in ['.pdf', '.docx']:
        return jsonify({'error': 'Unsupported file type. Please upload a PDF or Word document.'})

    file_id = str(uuid.uuid4())
    input_path = UPLOAD_FOLDER / f"{file_id}{ext}"
    file.save(str(input_path))

    detector = SensitiveInfoDetector()
    categories = request.form.getlist('categories')
    for cat_key in categories:
        if cat_key in PREDEFINED_CATEGORIES:
            detector.add_category(cat_key, PREDEFINED_CATEGORIES[cat_key])

    custom_terms = request.form.get('custom_terms', '').strip().split('\n')
    detector.set_custom_terms(custom_terms)

    try:
        text = extract_text(str(input_path))
        matches = detector.find_sensitive_text(text)

        preview_html = ""
        last_end = 0
        for start, end, matched_text, category in matches:
            preview_html += escape_html(text[last_end:start])
            preview_html += f'<span class="match">{escape_html(matched_text)}</span>'
            last_end = end
        preview_html += escape_html(text[last_end:])

        category_counts = {}
        for _, _, _, category in matches:
            category_counts[category] = category_counts.get(category, 0) + 1

        os.remove(input_path)

        return jsonify({
            'total': len(matches),
            'categories': category_counts,
            'preview_html': preview_html[:50000]
        })

    except Exception as e:
        if input_path.exists():
            os.remove(input_path)
        return jsonify({'error': str(e)})


@app.route('/redact-batch', methods=['POST'])
def redact_batch():
    if 'files' not in request.files:
        return jsonify({'error': 'No files uploaded'})

    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'No files selected'})

    # Setup detector
    detector = SensitiveInfoDetector()
    categories = request.form.getlist('categories')
    for cat_key in categories:
        if cat_key in PREDEFINED_CATEGORIES:
            detector.add_category(cat_key, PREDEFINED_CATEGORIES[cat_key])

    custom_terms = request.form.get('custom_terms', '').strip().split('\n')
    detector.set_custom_terms(custom_terms)

    # Process each file
    batch_id = str(uuid.uuid4())
    batch_folder = OUTPUT_FOLDER / batch_id
    batch_folder.mkdir(exist_ok=True)

    results = []
    total_redactions = 0
    total_categories = {}

    for file in files:
        ext = Path(file.filename).suffix.lower()
        if ext not in ['.pdf', '.docx']:
            continue

        file_id = str(uuid.uuid4())
        input_path = UPLOAD_FOLDER / f"{file_id}{ext}"
        file.save(str(input_path))

        original_name = Path(file.filename).stem
        output_filename = f"{original_name}_redacted{ext}"
        output_path = batch_folder / output_filename

        try:
            if ext == '.pdf':
                stats = redact_pdf(str(input_path), str(output_path), detector)
            else:
                stats = redact_docx(str(input_path), str(output_path), detector)

            results.append({
                'original_name': file.filename,
                'output_name': output_filename,
                'redactions': stats['redactions'],
                'categories': stats['categories']
            })

            total_redactions += stats['redactions']
            for cat, count in stats['categories'].items():
                total_categories[cat] = total_categories.get(cat, 0) + count

            os.remove(input_path)

        except Exception as e:
            if input_path.exists():
                os.remove(input_path)
            results.append({
                'original_name': file.filename,
                'error': str(e)
            })

    # If multiple files, create ZIP
    if len(results) > 1:
        zip_path = OUTPUT_FOLDER / f"{batch_id}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zf:
            for file in batch_folder.iterdir():
                zf.write(file, file.name)
        download_id = f"{batch_id}.zip"
    else:
        # Single file - get its path
        if results and 'output_name' in results[0]:
            download_id = f"{batch_id}/{results[0]['output_name']}"
        else:
            return jsonify({'error': 'Failed to process file'})

    return jsonify({
        'success': True,
        'download_id': download_id,
        'file_count': len(results),
        'total_redactions': total_redactions,
        'total_categories': total_categories,
        'files': results
    })


@app.route('/download/<path:download_id>')
def download(download_id):
    # Handle ZIP downloads
    if download_id.endswith('.zip'):
        zip_path = OUTPUT_FOLDER / download_id
        if zip_path.exists():
            return send_file(
                str(zip_path),
                as_attachment=True,
                download_name='redacted_documents.zip'
            )
    else:
        # Handle single file downloads
        file_path = OUTPUT_FOLDER / download_id
        if file_path.exists():
            return send_file(
                str(file_path),
                as_attachment=True,
                download_name=file_path.name
            )

    return "File not found", 404


def escape_html(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


@app.route('/health')
def health():
    """Health check endpoint for keep-alive pings"""
    return jsonify({'status': 'ok', 'timestamp': time.time()})


def keep_alive():
    """Background thread to ping the app every 10 minutes to prevent Render from sleeping"""
    app_url = os.environ.get('RENDER_EXTERNAL_URL')
    if not app_url:
        return  # Only run on Render

    health_url = f"{app_url}/health"
    while True:
        time.sleep(600)  # 10 minutes
        try:
            urllib.request.urlopen(health_url, timeout=30)
            print(f"Keep-alive ping sent to {health_url}")
        except Exception as e:
            print(f"Keep-alive ping failed: {e}")


# Start keep-alive thread on Render
if os.environ.get('RENDER_EXTERNAL_URL'):
    keep_alive_thread = threading.Thread(target=keep_alive, daemon=True)
    keep_alive_thread.start()
    print("Keep-alive thread started")


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    debug = os.environ.get('FLASK_ENV') != 'production'

    if debug:
        print("\n" + "="*50)
        print("Document Redaction Tool - Web Interface")
        print("="*50)
        print(f"\nOpen your browser and go to:")
        print(f"\n  http://localhost:{port}")
        print("\n" + "="*50 + "\n")

    app.run(debug=debug, host='0.0.0.0', port=port)
