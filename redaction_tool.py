#!/usr/bin/env python3
"""
Document Redaction Tool
Redacts sensitive information from PDFs and Word documents.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import shutil


@dataclass
class SensitiveCategory:
    """Represents a category of sensitive information to redact."""
    name: str
    description: str
    patterns: list = field(default_factory=list)
    examples: list = field(default_factory=list)
    enabled: bool = True
    use_exact_match: bool = False  # If True, only match exact examples


# Predefined categories with common patterns
PREDEFINED_CATEGORIES = {
    "SSN": SensitiveCategory(
        name="Social Security Numbers",
        description="US Social Security Numbers (XXX-XX-XXXX format)",
        patterns=[r'\b\d{3}-\d{2}-\d{4}\b', r'\b\d{9}\b'],
        examples=["123-45-6789"]
    ),
    "Email": SensitiveCategory(
        name="Email Addresses",
        description="Email addresses",
        patterns=[r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'],
        examples=["example@email.com"]
    ),
    "Phone": SensitiveCategory(
        name="Phone Numbers",
        description="US phone numbers in various formats",
        patterns=[
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}\b'
        ],
        examples=["555-123-4567", "(555) 123-4567"]
    ),
    "CreditCard": SensitiveCategory(
        name="Credit Card Numbers",
        description="Credit card numbers (13-19 digits)",
        patterns=[r'\b(?:\d{4}[-\s]?){3,4}\d{1,4}\b'],
        examples=["4111-1111-1111-1111"]
    ),
    "Date": SensitiveCategory(
        name="Dates",
        description="Dates in various formats",
        patterns=[
            r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',
            r'\b\d{1,2}-\d{1,2}-\d{2,4}\b',
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b'
        ],
        examples=["01/15/2024", "January 15, 2024"]
    ),
    "Custom": SensitiveCategory(
        name="Custom Terms",
        description="Custom words, phrases, or patterns you define",
        patterns=[],
        examples=[],
        use_exact_match=True
    )
}


class SensitiveInfoDetector:
    """Detects sensitive information in text based on configured categories."""

    def __init__(self):
        self.categories = {}

    def add_category(self, key: str, category: SensitiveCategory):
        """Add a category of sensitive information to detect."""
        self.categories[key] = category

    def remove_category(self, key: str):
        """Remove a category."""
        if key in self.categories:
            del self.categories[key]

    def find_sensitive_text(self, text: str) -> list:
        """
        Find all sensitive text matches in the given text.
        Returns list of (start, end, matched_text, category_name) tuples.
        """
        matches = []

        for key, category in self.categories.items():
            if not category.enabled:
                continue

            # Match against patterns
            for pattern in category.patterns:
                try:
                    for match in re.finditer(pattern, text, re.IGNORECASE):
                        matches.append((
                            match.start(),
                            match.end(),
                            match.group(),
                            category.name
                        ))
                except re.error:
                    continue

            # Match exact examples if use_exact_match is True
            if category.use_exact_match:
                for example in category.examples:
                    if example.strip():
                        # Escape special regex characters for literal matching
                        escaped = re.escape(example)
                        try:
                            for match in re.finditer(escaped, text, re.IGNORECASE):
                                matches.append((
                                    match.start(),
                                    match.end(),
                                    match.group(),
                                    category.name
                                ))
                        except re.error:
                            continue

        # Remove overlapping matches (keep longer ones)
        matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))
        filtered = []
        last_end = -1
        for match in matches:
            if match[0] >= last_end:
                filtered.append(match)
                last_end = match[1]

        return filtered


class PDFRedactor:
    """Handles PDF document redaction."""

    def __init__(self, detector: SensitiveInfoDetector):
        self.detector = detector

    def redact(self, input_path: str, output_path: str) -> dict:
        """
        Redact sensitive information from a PDF.
        Returns statistics about redactions made.
        """
        doc = fitz.open(input_path)
        stats = {"pages": 0, "redactions": 0, "categories": {}}

        for page_num, page in enumerate(doc):
            stats["pages"] += 1
            text_dict = page.get_text("dict")

            # Process each text block
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:  # Not a text block
                    continue

                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text:
                            continue

                        matches = self.detector.find_sensitive_text(text)

                        for start, end, matched_text, category in matches:
                            # Find the position of this text on the page
                            instances = page.search_for(matched_text)

                            for inst in instances:
                                # Add redaction annotation
                                page.add_redact_annot(inst, fill=(0, 0, 0))
                                stats["redactions"] += 1
                                stats["categories"][category] = stats["categories"].get(category, 0) + 1

            # Apply all redactions on this page
            page.apply_redactions()

        # Save the redacted document
        doc.save(output_path)
        doc.close()

        return stats


class WordRedactor:
    """Handles Word document redaction."""

    def __init__(self, detector: SensitiveInfoDetector):
        self.detector = detector

    def _create_redacted_run(self, text: str, original_run):
        """Create a redacted version of text (black boxes represented as highlighted black text)."""
        # Return black squares to represent redaction
        return "█" * len(text)

    def redact(self, input_path: str, output_path: str) -> dict:
        """
        Redact sensitive information from a Word document.
        Returns statistics about redactions made.
        """
        doc = Document(input_path)
        stats = {"paragraphs": 0, "redactions": 0, "categories": {}}

        # Process paragraphs
        for para in doc.paragraphs:
            stats["paragraphs"] += 1
            self._process_paragraph(para, stats)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        stats["paragraphs"] += 1
                        self._process_paragraph(para, stats)

        # Process headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        stats["paragraphs"] += 1
                        self._process_paragraph(para, stats)

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        stats["paragraphs"] += 1
                        self._process_paragraph(para, stats)

        doc.save(output_path)
        return stats

    def _process_paragraph(self, para, stats: dict):
        """Process a single paragraph for redaction."""
        full_text = para.text
        if not full_text:
            return

        matches = self.detector.find_sensitive_text(full_text)
        if not matches:
            return

        # Build new text with redactions
        new_text = list(full_text)
        for start, end, matched_text, category in matches:
            for i in range(start, end):
                new_text[i] = "█"
            stats["redactions"] += 1
            stats["categories"][category] = stats["categories"].get(category, 0) + 1

        new_text = "".join(new_text)

        # Clear existing runs and add new text
        # Preserve formatting from first run if available
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic

            # Clear all runs
            for run in para.runs:
                run.text = ""

            # Set new text on first run
            para.runs[0].text = new_text
            para.runs[0].font.name = font_name
            para.runs[0].font.size = font_size
            para.runs[0].font.bold = bold
            para.runs[0].font.italic = italic


class RedactionToolGUI:
    """Main GUI application for the redaction tool."""

    def __init__(self, root):
        self.root = root
        self.root.title("Document Redaction Tool")
        self.root.geometry("900x700")
        self.root.minsize(800, 600)

        self.detector = SensitiveInfoDetector()
        self.loaded_file = None
        self.category_vars = {}

        self._setup_styles()
        self._create_widgets()
        self._initialize_categories()

    def _setup_styles(self):
        """Configure ttk styles."""
        style = ttk.Style()
        style.configure("Header.TLabel", font=("Helvetica", 12, "bold"))
        style.configure("Status.TLabel", font=("Helvetica", 10))

    def _create_widgets(self):
        """Create all GUI widgets."""
        # Main container with padding
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        # File selection section
        file_frame = ttk.LabelFrame(main_frame, text="Document Selection", padding="10")
        file_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 10))

        self.file_label = ttk.Label(file_frame, text="No file selected", style="Status.TLabel")
        self.file_label.grid(row=0, column=0, sticky="w", padx=(0, 10))

        ttk.Button(file_frame, text="Browse...", command=self._browse_file).grid(row=0, column=1)
        ttk.Button(file_frame, text="Clear", command=self._clear_file).grid(row=0, column=2, padx=(5, 0))

        file_frame.columnconfigure(0, weight=1)

        # Categories section
        cat_frame = ttk.LabelFrame(main_frame, text="Sensitive Information Categories", padding="10")
        cat_frame.grid(row=1, column=0, sticky="nsew", pady=(0, 10), padx=(0, 5))

        # Checkboxes for predefined categories
        ttk.Label(cat_frame, text="Select categories to redact:", style="Header.TLabel").grid(
            row=0, column=0, columnspan=2, sticky="w", pady=(0, 10))

        self.category_checkboxes = {}
        row = 1
        for key, cat in PREDEFINED_CATEGORIES.items():
            if key != "Custom":
                var = tk.BooleanVar(value=True)
                self.category_vars[key] = var
                cb = ttk.Checkbutton(cat_frame, text=cat.name, variable=var,
                                     command=lambda k=key: self._toggle_category(k))
                cb.grid(row=row, column=0, sticky="w", pady=2)
                self.category_checkboxes[key] = cb

                # Example label
                example_text = f"e.g., {cat.examples[0]}" if cat.examples else ""
                ttk.Label(cat_frame, text=example_text, foreground="gray").grid(
                    row=row, column=1, sticky="w", padx=(10, 0))
                row += 1

        # Custom category checkbox
        self.custom_var = tk.BooleanVar(value=False)
        self.category_vars["Custom"] = self.custom_var
        ttk.Checkbutton(cat_frame, text="Custom Terms", variable=self.custom_var,
                       command=lambda: self._toggle_category("Custom")).grid(
            row=row, column=0, sticky="w", pady=2)

        cat_frame.columnconfigure(1, weight=1)

        # Custom terms section
        custom_frame = ttk.LabelFrame(main_frame, text="Custom Terms & Patterns", padding="10")
        custom_frame.grid(row=1, column=1, sticky="nsew", pady=(0, 10), padx=(5, 0))

        ttk.Label(custom_frame, text="Enter terms to redact (one per line):",
                 style="Header.TLabel").grid(row=0, column=0, sticky="w", pady=(0, 5))

        self.custom_text = scrolledtext.ScrolledText(custom_frame, width=30, height=10,
                                                      font=("Courier", 10))
        self.custom_text.grid(row=1, column=0, sticky="nsew", pady=(0, 10))

        ttk.Label(custom_frame, text="Examples:\nJohn Doe\nProject Alpha\nConfidential",
                 foreground="gray").grid(row=2, column=0, sticky="w")

        ttk.Button(custom_frame, text="Update Custom Terms",
                  command=self._update_custom_terms).grid(row=3, column=0, pady=(10, 0))

        custom_frame.columnconfigure(0, weight=1)
        custom_frame.rowconfigure(1, weight=1)

        # Preview section
        preview_frame = ttk.LabelFrame(main_frame, text="Preview & Results", padding="10")
        preview_frame.grid(row=2, column=0, columnspan=2, sticky="nsew", pady=(0, 10))

        self.preview_text = scrolledtext.ScrolledText(preview_frame, width=80, height=12,
                                                       font=("Courier", 10), state="disabled")
        self.preview_text.grid(row=0, column=0, sticky="nsew")

        # Configure tag for highlighting matches
        self.preview_text.tag_configure("match", background="yellow")
        self.preview_text.tag_configure("redacted", background="black", foreground="black")

        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(0, weight=1)

        # Action buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=2, sticky="ew")

        ttk.Button(button_frame, text="Preview Redactions",
                  command=self._preview_redactions).grid(row=0, column=0, padx=(0, 10))
        ttk.Button(button_frame, text="Redact Document",
                  command=self._redact_document).grid(row=0, column=1, padx=(0, 10))

        self.status_label = ttk.Label(button_frame, text="Ready", style="Status.TLabel")
        self.status_label.grid(row=0, column=2, sticky="e")

        button_frame.columnconfigure(2, weight=1)

        # Configure grid weights
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        main_frame.rowconfigure(2, weight=2)

    def _initialize_categories(self):
        """Initialize the detector with predefined categories."""
        for key, cat in PREDEFINED_CATEGORIES.items():
            if key != "Custom":
                self.detector.add_category(key, cat)

    def _toggle_category(self, key: str):
        """Toggle a category on/off."""
        enabled = self.category_vars[key].get()

        if key == "Custom":
            if enabled:
                self._update_custom_terms()
            else:
                self.detector.remove_category("Custom")
        else:
            if key in self.detector.categories:
                self.detector.categories[key].enabled = enabled

    def _update_custom_terms(self):
        """Update the custom terms category."""
        terms = self.custom_text.get("1.0", tk.END).strip().split("\n")
        terms = [t.strip() for t in terms if t.strip()]

        custom_cat = SensitiveCategory(
            name="Custom Terms",
            description="User-defined terms",
            patterns=[],
            examples=terms,
            enabled=self.custom_var.get(),
            use_exact_match=True
        )

        self.detector.add_category("Custom", custom_cat)
        self.status_label.config(text=f"Updated {len(terms)} custom terms")

    def _browse_file(self):
        """Open file browser dialog."""
        filetypes = [
            ("Supported Documents", "*.pdf *.docx"),
            ("PDF Files", "*.pdf"),
            ("Word Documents", "*.docx"),
            ("All Files", "*.*")
        ]

        filepath = filedialog.askopenfilename(
            title="Select Document to Redact",
            filetypes=filetypes
        )

        if filepath:
            self.loaded_file = filepath
            filename = os.path.basename(filepath)
            self.file_label.config(text=f"Selected: {filename}")
            self.status_label.config(text=f"Loaded: {filename}")

    def _clear_file(self):
        """Clear the selected file."""
        self.loaded_file = None
        self.file_label.config(text="No file selected")
        self.status_label.config(text="Ready")
        self._clear_preview()

    def _clear_preview(self):
        """Clear the preview text."""
        self.preview_text.config(state="normal")
        self.preview_text.delete("1.0", tk.END)
        self.preview_text.config(state="disabled")

    def _extract_text(self, filepath: str) -> str:
        """Extract text from a document for preview."""
        ext = Path(filepath).suffix.lower()

        if ext == ".pdf":
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text

        elif ext == ".docx":
            doc = Document(filepath)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text

        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _preview_redactions(self):
        """Preview what will be redacted."""
        if not self.loaded_file:
            messagebox.showwarning("No File", "Please select a document first.")
            return

        try:
            self.status_label.config(text="Analyzing document...")
            self.root.update()

            text = self._extract_text(self.loaded_file)
            matches = self.detector.find_sensitive_text(text)

            self.preview_text.config(state="normal")
            self.preview_text.delete("1.0", tk.END)

            # Show preview with highlighted matches
            last_end = 0
            for start, end, matched_text, category in matches:
                # Insert text before match
                self.preview_text.insert(tk.END, text[last_end:start])
                # Insert matched text with highlight
                self.preview_text.insert(tk.END, matched_text, "match")
                last_end = end

            # Insert remaining text
            self.preview_text.insert(tk.END, text[last_end:])

            self.preview_text.config(state="disabled")

            # Update status
            category_counts = {}
            for _, _, _, category in matches:
                category_counts[category] = category_counts.get(category, 0) + 1

            status_parts = [f"{count} {cat}" for cat, count in category_counts.items()]
            status = f"Found {len(matches)} items to redact: " + ", ".join(status_parts) if matches else "No sensitive information found"
            self.status_label.config(text=status)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to analyze document: {str(e)}")
            self.status_label.config(text="Error analyzing document")

    def _redact_document(self):
        """Perform the actual redaction."""
        if not self.loaded_file:
            messagebox.showwarning("No File", "Please select a document first.")
            return

        # Get save location
        ext = Path(self.loaded_file).suffix.lower()
        default_name = Path(self.loaded_file).stem + "_redacted" + ext

        filetypes = [("PDF Files", "*.pdf")] if ext == ".pdf" else [("Word Documents", "*.docx")]

        output_path = filedialog.asksaveasfilename(
            title="Save Redacted Document",
            defaultextension=ext,
            initialfile=default_name,
            filetypes=filetypes
        )

        if not output_path:
            return

        try:
            self.status_label.config(text="Redacting document...")
            self.root.update()

            if ext == ".pdf":
                redactor = PDFRedactor(self.detector)
            else:
                redactor = WordRedactor(self.detector)

            stats = redactor.redact(self.loaded_file, output_path)

            # Show results
            self.preview_text.config(state="normal")
            self.preview_text.delete("1.0", tk.END)

            result_text = f"Redaction Complete!\n\n"
            result_text += f"Output file: {output_path}\n\n"
            result_text += f"Statistics:\n"
            result_text += f"  Total redactions: {stats['redactions']}\n"

            if stats['categories']:
                result_text += f"\n  By category:\n"
                for cat, count in stats['categories'].items():
                    result_text += f"    - {cat}: {count}\n"

            self.preview_text.insert(tk.END, result_text)
            self.preview_text.config(state="disabled")

            self.status_label.config(text=f"Redacted {stats['redactions']} items")

            messagebox.showinfo("Success",
                f"Document redacted successfully!\n\n"
                f"Redactions made: {stats['redactions']}\n"
                f"Saved to: {output_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to redact document: {str(e)}")
            self.status_label.config(text="Error during redaction")


def main():
    """Main entry point."""
    root = tk.Tk()
    app = RedactionToolGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
