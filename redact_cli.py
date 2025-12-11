#!/usr/bin/env python3
"""
Document Redaction Tool - Command Line Interface
Redacts sensitive information from PDFs and Word documents.
"""

import argparse
import os
import re
import sys
from pathlib import Path
from dataclasses import dataclass, field


@dataclass
class SensitiveCategory:
    """Represents a category of sensitive information to redact."""
    name: str
    description: str
    patterns: list = field(default_factory=list)
    examples: list = field(default_factory=list)
    enabled: bool = True
    use_exact_match: bool = False


# Predefined categories with common patterns
PREDEFINED_CATEGORIES = {
    "ssn": SensitiveCategory(
        name="Social Security Numbers",
        description="US Social Security Numbers (XXX-XX-XXXX format)",
        patterns=[r'\b\d{3}-\d{2}-\d{4}\b', r'\b\d{9}\b'],
        examples=["123-45-6789"]
    ),
    "email": SensitiveCategory(
        name="Email Addresses",
        description="Email addresses",
        patterns=[r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'],
        examples=["example@email.com"]
    ),
    "phone": SensitiveCategory(
        name="Phone Numbers",
        description="US phone numbers in various formats",
        patterns=[
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}\b'
        ],
        examples=["555-123-4567", "(555) 123-4567"]
    ),
    "creditcard": SensitiveCategory(
        name="Credit Card Numbers",
        description="Credit card numbers (13-19 digits)",
        patterns=[r'\b(?:\d{4}[-\s]?){3,4}\d{1,4}\b'],
        examples=["4111-1111-1111-1111"]
    ),
    "date": SensitiveCategory(
        name="Dates",
        description="Dates in various formats",
        patterns=[
            r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',
            r'\b\d{1,2}-\d{1,2}-\d{2,4}\b',
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b'
        ],
        examples=["01/15/2024", "January 15, 2024"]
    ),
}


class SensitiveInfoDetector:
    """Detects sensitive information in text based on configured categories."""

    def __init__(self):
        self.categories = {}
        self.custom_terms = []

    def add_category(self, key: str, category: SensitiveCategory):
        self.categories[key] = category

    def set_custom_terms(self, terms: list):
        self.custom_terms = [t.strip() for t in terms if t.strip()]

    def find_sensitive_text(self, text: str) -> list:
        """Find all sensitive text matches. Returns list of (start, end, matched_text, category_name)."""
        matches = []

        for key, category in self.categories.items():
            if not category.enabled:
                continue

            for pattern in category.patterns:
                try:
                    for match in re.finditer(pattern, text, re.IGNORECASE):
                        matches.append((match.start(), match.end(), match.group(), category.name))
                except re.error:
                    continue

        # Custom terms (exact match)
        for term in self.custom_terms:
            escaped = re.escape(term)
            try:
                for match in re.finditer(escaped, text, re.IGNORECASE):
                    matches.append((match.start(), match.end(), match.group(), "Custom Terms"))
            except re.error:
                continue

        # Remove overlapping matches
        matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))
        filtered = []
        last_end = -1
        for match in matches:
            if match[0] >= last_end:
                filtered.append(match)
                last_end = match[1]

        return filtered


def redact_pdf(input_path: str, output_path: str, detector: SensitiveInfoDetector) -> dict:
    """Redact sensitive information from a PDF."""
    import fitz

    doc = fitz.open(input_path)
    stats = {"pages": 0, "redactions": 0, "categories": {}}

    for page in doc:
        stats["pages"] += 1
        text_dict = page.get_text("dict")

        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue

            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text:
                        continue

                    matches = detector.find_sensitive_text(text)
                    for start, end, matched_text, category in matches:
                        instances = page.search_for(matched_text)
                        for inst in instances:
                            page.add_redact_annot(inst, fill=(0, 0, 0))
                            stats["redactions"] += 1
                            stats["categories"][category] = stats["categories"].get(category, 0) + 1

        page.apply_redactions()

    doc.save(output_path)
    doc.close()
    return stats


def redact_docx(input_path: str, output_path: str, detector: SensitiveInfoDetector) -> dict:
    """Redact sensitive information from a Word document."""
    from docx import Document

    doc = Document(input_path)
    stats = {"paragraphs": 0, "redactions": 0, "categories": {}}

    def process_paragraph(para):
        full_text = para.text
        if not full_text:
            return

        matches = detector.find_sensitive_text(full_text)
        if not matches:
            return

        new_text = list(full_text)
        for start, end, matched_text, category in matches:
            for i in range(start, end):
                new_text[i] = "█"
            stats["redactions"] += 1
            stats["categories"][category] = stats["categories"].get(category, 0) + 1

        new_text = "".join(new_text)

        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic

            for run in para.runs:
                run.text = ""

            para.runs[0].text = new_text
            para.runs[0].font.name = font_name
            para.runs[0].font.size = font_size
            para.runs[0].font.bold = bold
            para.runs[0].font.italic = italic

    for para in doc.paragraphs:
        stats["paragraphs"] += 1
        process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    stats["paragraphs"] += 1
                    process_paragraph(para)

    doc.save(output_path)
    return stats


def preview_document(file_path: str, detector: SensitiveInfoDetector):
    """Preview what would be redacted."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        import fitz
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
    elif ext == ".docx":
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        print(f"Error: Unsupported file type: {ext}")
        sys.exit(1)

    matches = detector.find_sensitive_text(text)

    print("\n" + "="*60)
    print("PREVIEW: Items that will be redacted")
    print("="*60 + "\n")

    if not matches:
        print("No sensitive information found.")
        return

    # Group by category
    by_category = {}
    for start, end, matched_text, category in matches:
        if category not in by_category:
            by_category[category] = []
        by_category[category].append(matched_text)

    for category, items in by_category.items():
        print(f"\n{category}:")
        for item in items:
            print(f"  • {item}")

    print(f"\n{'='*60}")
    print(f"Total: {len(matches)} items to redact")
    print("="*60)


def main():
    parser = argparse.ArgumentParser(
        description="Redact sensitive information from PDF and Word documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Redact all default categories (SSN, email, phone, credit card, dates)
  python3 redact_cli.py document.pdf

  # Redact only emails and phone numbers
  python3 redact_cli.py document.docx --categories email phone

  # Add custom terms to redact
  python3 redact_cli.py document.pdf --custom "John Doe" "Project Alpha" "Confidential"

  # Preview what would be redacted (no changes made)
  python3 redact_cli.py document.pdf --preview

  # Specify output file name
  python3 redact_cli.py document.pdf -o redacted_output.pdf

Available categories: ssn, email, phone, creditcard, date
        """
    )

    parser.add_argument("input", help="Input PDF or Word document")
    parser.add_argument("-o", "--output", help="Output file path (default: input_redacted.ext)")
    parser.add_argument(
        "-c", "--categories",
        nargs="+",
        choices=list(PREDEFINED_CATEGORIES.keys()) + ["all"],
        default=["all"],
        help="Categories to redact (default: all)"
    )
    parser.add_argument(
        "--custom",
        nargs="+",
        default=[],
        help="Custom terms/names to redact"
    )
    parser.add_argument(
        "--preview",
        action="store_true",
        help="Preview what would be redacted without making changes"
    )

    args = parser.parse_args()

    # Validate input file
    if not os.path.exists(args.input):
        print(f"Error: File not found: {args.input}")
        sys.exit(1)

    ext = Path(args.input).suffix.lower()
    if ext not in [".pdf", ".docx"]:
        print(f"Error: Unsupported file type: {ext}")
        print("Supported formats: .pdf, .docx")
        sys.exit(1)

    # Setup detector
    detector = SensitiveInfoDetector()

    categories_to_use = args.categories
    if "all" in categories_to_use:
        categories_to_use = list(PREDEFINED_CATEGORIES.keys())

    for cat_key in categories_to_use:
        detector.add_category(cat_key, PREDEFINED_CATEGORIES[cat_key])

    if args.custom:
        detector.set_custom_terms(args.custom)

    print(f"\nProcessing: {args.input}")
    print(f"Categories: {', '.join(categories_to_use)}")
    if args.custom:
        print(f"Custom terms: {', '.join(args.custom)}")

    # Preview mode
    if args.preview:
        preview_document(args.input, detector)
        return

    # Determine output path
    if args.output:
        output_path = args.output
    else:
        stem = Path(args.input).stem
        output_path = str(Path(args.input).parent / f"{stem}_redacted{ext}")

    # Perform redaction
    print(f"\nRedacting document...")

    try:
        if ext == ".pdf":
            stats = redact_pdf(args.input, output_path, detector)
        else:
            stats = redact_docx(args.input, output_path, detector)

        print(f"\n{'='*60}")
        print("REDACTION COMPLETE")
        print("="*60)
        print(f"\nOutput file: {output_path}")
        print(f"Total redactions: {stats['redactions']}")

        if stats['categories']:
            print("\nBy category:")
            for cat, count in stats['categories'].items():
                print(f"  • {cat}: {count}")

        print("="*60 + "\n")

    except Exception as e:
        print(f"\nError: Failed to redact document: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
